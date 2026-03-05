"""
PDF Failure Analysis Lambda

Analyzes PDFs that fail during Adobe API processing to identify
likely causes of failure (excluding rate limit errors).
"""

import json
import os
import logging
import boto3
from datetime import datetime, timezone
from io import BytesIO
from analyzer import analyze_pdf, AnalysisResult

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

s3 = boto3.client('s3')
dynamodb = boto3.resource('dynamodb')

# Environment variables
REPORT_BUCKET = os.environ.get('REPORT_BUCKET', '')
SAVE_REPORTS_TO_S3 = os.environ.get('SAVE_REPORTS_TO_S3', 'true').lower() == 'true'
ANALYSIS_TABLE = os.environ.get('ANALYSIS_TABLE', '')


def format_text_report(log_entry: dict, result: AnalysisResult) -> str:
    """Format the analysis result as a human-readable text report."""
    lines = []
    lines.append("=" * 70)
    lines.append("PDF FAILURE ANALYSIS REPORT")
    lines.append("=" * 70)
    lines.append("")
    lines.append(f"Filename:       {log_entry['filename']}")
    lines.append(f"S3 Location:    s3://{log_entry['s3_bucket']}/{log_entry['s3_key']}")
    lines.append(f"API Type:       {log_entry['api_type']}")
    lines.append(f"Analysis Time:  {log_entry['timestamp']}")
    lines.append("")
    lines.append("-" * 70)
    lines.append("ORIGINAL ERROR")
    lines.append("-" * 70)
    lines.append(log_entry['original_error'])
    lines.append("")
    lines.append("-" * 70)
    lines.append("PDF PROPERTIES")
    lines.append("-" * 70)
    lines.append(f"File Size:      {result.file_size_mb} MB")
    lines.append(f"Page Count:     {result.page_count}")
    lines.append(f"Image Count:    {result.image_count}")
    lines.append(f"Font Count:     {result.font_count}")
    lines.append(f"Encrypted:      {'Yes' if result.has_encryption else 'No'}")
    lines.append(f"PDF Version:    {result.pdf_version}")
    lines.append("")
    
    if result.issues:
        lines.append("-" * 70)
        lines.append("ISSUES FOUND")
        lines.append("-" * 70)
        for i, issue in enumerate(result.issues, 1):
            lines.append(f"\n{i}. [{issue.severity.value}] {issue.category.value}")
            lines.append(f"   {issue.description}")
            if issue.details:
                for detail in issue.details[:5]:
                    lines.append(f"   - {detail}")
                if len(issue.details) > 5:
                    lines.append(f"   ... and {len(issue.details) - 5} more")
    else:
        lines.append("-" * 70)
        lines.append("No structural issues detected")
        lines.append("-" * 70)
    
    lines.append("")
    lines.append("=" * 70)
    lines.append("LIKELY CAUSE")
    lines.append("=" * 70)
    lines.append(result.likely_cause or "Unknown")
    lines.append("")
    
    if result.analysis_error:
        lines.append("-" * 70)
        lines.append("ANALYSIS ERROR")
        lines.append("-" * 70)
        lines.append(result.analysis_error)
    
    return "\n".join(lines)


def create_docx_report(log_entry: dict, result: AnalysisResult) -> bytes:
    """Create a formatted Word document report."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Title
    title = doc.add_heading('PDF Failure Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # File info section
    doc.add_heading('File Information', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    rows_data = [
        ('Filename', log_entry['filename']),
        ('S3 Location', f"s3://{log_entry['s3_bucket']}/{log_entry['s3_key']}"),
        ('API Type', log_entry['api_type']),
        ('Analysis Time', log_entry['timestamp']),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    
    # Original error
    doc.add_heading('Original Error', level=1)
    error_para = doc.add_paragraph(log_entry['original_error'])
    error_para.style = 'Quote'
    
    # PDF Properties
    doc.add_heading('PDF Properties', level=1)
    props_table = doc.add_table(rows=6, cols=2)
    props_table.style = 'Table Grid'
    
    props_data = [
        ('File Size', f"{result.file_size_mb} MB"),
        ('Page Count', str(result.page_count)),
        ('Image Count', str(result.image_count)),
        ('Font Count', str(result.font_count)),
        ('Encrypted', 'Yes' if result.has_encryption else 'No'),
        ('PDF Version', result.pdf_version),
    ]
    for i, (label, value) in enumerate(props_data):
        props_table.rows[i].cells[0].text = label
        props_table.rows[i].cells[1].text = value
    
    # Issues Found
    doc.add_heading('Issues Found', level=1)
    if result.issues:
        for issue in result.issues:
            # Issue header
            p = doc.add_paragraph()
            severity_run = p.add_run(f"[{issue.severity.value}] ")
            severity_run.bold = True
            if issue.severity.value == 'HIGH':
                severity_run.font.color.rgb = None  # Would need python-docx color support
            p.add_run(f"{issue.category.value}")
            
            # Description
            doc.add_paragraph(issue.description, style='List Bullet')
            
            # Details
            if issue.details:
                for detail in issue.details[:5]:
                    doc.add_paragraph(detail, style='List Bullet 2')
                if len(issue.details) > 5:
                    doc.add_paragraph(f"... and {len(issue.details) - 5} more", style='List Bullet 2')
    else:
        doc.add_paragraph('No structural issues detected.')
    
    # Likely Cause
    doc.add_heading('Likely Cause', level=1)
    cause_para = doc.add_paragraph(result.likely_cause or 'Unknown')
    cause_para.style = 'Intense Quote'
    
    # Analysis error if any
    if result.analysis_error:
        doc.add_heading('Analysis Error', level=1)
        doc.add_paragraph(result.analysis_error)
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def save_to_dynamodb(log_entry: dict, result: AnalysisResult, started_at: str = None, failed_at: str = None) -> bool:
    """Save analysis data to DynamoDB table.
    
    Timing data (started_at, failed_at) is stored permanently so it's available
    even after the in-flight tracker entries expire.
    """
    if not ANALYSIS_TABLE:
        logger.warning("ANALYSIS_TABLE not configured, skipping DynamoDB save")
        return False
    
    try:
        table = dynamodb.Table(ANALYSIS_TABLE)
        
        # Build the item with all analysis data
        item = {
            's3_key': log_entry['s3_key'],
            'analysis_timestamp': log_entry['timestamp'],
            'filename': log_entry['filename'],
            's3_bucket': log_entry['s3_bucket'],
            'api_type': log_entry['api_type'],
            'original_error': log_entry['original_error'],
            'file_size_mb': str(result.file_size_mb),
            'page_count': result.page_count,
            'image_count': result.image_count,
            'font_count': result.font_count,
            'has_encryption': result.has_encryption,
            'pdf_version': result.pdf_version,
            'likely_cause': result.likely_cause or 'Unknown',
            'issue_count': len(result.issues),
            'issues': json.dumps([{
                'severity': issue.severity.value,
                'category': issue.category.value,
                'description': issue.description,
                'details': issue.details
            } for issue in result.issues]),
            'analysis_error': result.analysis_error or ''
        }
        
        # Add timing data if provided (stored permanently for failure analysis reports)
        if started_at:
            item['started_at'] = started_at
        if failed_at:
            item['failed_at'] = failed_at
        
        table.put_item(Item=item)
        logger.info(f"Saved analysis data to DynamoDB: {log_entry['s3_key']} (started_at: {started_at}, failed_at: {failed_at})")
        return True
        
    except Exception as e:
        logger.error(f"Failed to save to DynamoDB: {e}")
        return False


def lambda_handler(event, context):
    """
    Analyze a PDF that failed during Adobe API processing.
    
    Expected event format:
    {
        "bucket": "source-bucket-name",
        "key": "path/to/file.pdf",
        "filename": "file.pdf",
        "original_error": "ServiceApiException: ...",
        "api_type": "autotag" | "extract",
        "started_at": "2026-03-05T15:00:00+00:00",  # Optional: when API call started
        "failed_at": "2026-03-05T15:02:39+00:00"    # Optional: when failure occurred
    }
    
    Can also be triggered via SNS/EventBridge with the payload in the message body.
    """
    logger.info(f"Received event: {json.dumps(event)}")
    
    # Handle SNS wrapper
    if 'Records' in event and event['Records']:
        record = event['Records'][0]
        if 'Sns' in record:
            event = json.loads(record['Sns']['Message'])
        elif 'body' in record:
            event = json.loads(record['body'])
    
    # Extract parameters
    bucket = event.get('bucket')
    key = event.get('key')
    filename = event.get('filename', os.path.basename(key) if key else 'unknown.pdf')
    original_error = event.get('original_error', 'Unknown error')
    api_type = event.get('api_type', 'unknown')
    started_at = event.get('started_at')  # Timing data from container
    failed_at = event.get('failed_at')    # Timing data from container
    
    if not bucket or not key:
        logger.error("Missing required parameters: bucket and key")
        return {
            'statusCode': 400,
            'body': json.dumps({'error': 'Missing required parameters: bucket and key'})
        }
    
    # Skip rate limit errors
    if '429' in original_error or 'Too Many Requests' in original_error:
        logger.info(f"Skipping analysis for rate limit error: {filename}")
        return {
            'statusCode': 200,
            'body': json.dumps({'message': 'Skipped - rate limit error'})
        }
    
    # Download PDF to /tmp
    local_path = f"/tmp/{filename}"
    try:
        logger.info(f"Downloading s3://{bucket}/{key} to {local_path}")
        s3.download_file(bucket, key, local_path)
    except Exception as e:
        logger.error(f"Failed to download PDF: {e}")
        return {
            'statusCode': 500,
            'body': json.dumps({'error': f'Failed to download PDF: {str(e)}'})
        }
    
    # Analyze the PDF
    try:
        logger.info(f"Analyzing PDF: {filename}")
        result = analyze_pdf(local_path)
    except Exception as e:
        logger.error(f"Analysis failed: {e}")
        result = AnalysisResult(
            filename=filename,
            file_size_mb=0,
            page_count=0,
            image_count=0,
            font_count=0,
            has_encryption=False,
            pdf_version="unknown",
            analysis_error=str(e),
            likely_cause="Analysis failed - PDF may be severely corrupted"
        )
    finally:
        # Clean up
        if os.path.exists(local_path):
            os.remove(local_path)
    
    # Build the log entry
    log_entry = {
        'event_type': 'PDF_FAILURE_ANALYSIS',
        'filename': filename,
        's3_bucket': bucket,
        's3_key': key,
        'api_type': api_type,
        'original_error': original_error,
        'analysis': result.to_dict(),
        'timestamp': datetime.now(timezone.utc).isoformat()
    }
    
    # Log structured output
    logger.info(f"PDF_FAILURE_ANALYSIS: {json.dumps(log_entry)}")
    
    # Save Word document report to S3
    if SAVE_REPORTS_TO_S3 and REPORT_BUCKET:
        timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = os.path.splitext(filename)[0]
        
        # Save Word document report only (no .txt version)
        try:
            docx_bytes = create_docx_report(log_entry, result)
            docx_key = f"reports/failure_analysis/{base_filename}_analysis_{timestamp_str}.docx"
            s3.put_object(
                Bucket=REPORT_BUCKET,
                Key=docx_key,
                Body=docx_bytes,
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            logger.info(f"Saved Word report to s3://{REPORT_BUCKET}/{docx_key}")
        except Exception as e:
            logger.warning(f"Failed to save Word report to S3: {e}")
    
    # Save analysis data to DynamoDB (includes timing data for permanent storage)
    save_to_dynamodb(log_entry, result, started_at=started_at, failed_at=failed_at)
    
    return {
        'statusCode': 200,
        'body': json.dumps({
            'filename': filename,
            'issues_found': len(result.issues),
            'likely_cause': result.likely_cause
        })
    }
